# -*- coding:utf-8 -*-
from bs4 import BeautifulSoup
import re
import json
import os
import copy
from docx import Document

htmlPath = 'C:\\Users\\Administrator\\Desktop\\新建文件夹\\非涉密终端检查报告_国泰新点_平台研发2部_曹政楠_20181221082509_1.html'

# 删除存放未处理文件的文本文档
untreatedFilesPath = "./untreatedFiles.txt"
if (os.path.exists(untreatedFilesPath)):
    os.remove(untreatedFilesPath)


# 删除文件
def deleteFile(path):
    try:
        os.remove(path)
        print('删除文件成功')
    except Exception as e:
        print('删除文件失败')
        print(e)
        untreatedFile(path)


# 修改并删除文件
def modifyAndDeleteFile(filePath, words, fileName, fileCategory):
    if (filePath.count('|') > 0):  # 压缩包不处理
        untreatedFile(filePath)
        return

    if (fileCategory == '文本文档' or fileCategory == 'XML 文件' or fileCategory == 'RTF 文件' or
            fileCategory == 'HTML 文档' or fileCategory == 'HTM 文件'):
        modifyTxt(filePath, words, fileName)
        return

    if (fileCategory == 'DOC 文档' or fileCategory == 'DOCX 文档'):
        modifyDoc(filePath, words, fileName)
        return
    untreatedFile(filePath)


# 记录未处理的文件
def untreatedFile(filePath):
    print('未处理文件')
    with open("./untreatedFiles.txt", "a") as f:
        f.write(filePath + '\r')


# 处理txt文件
def modifyTxt(filePath, words, fileName):
    wordlist = list(words)
    try:
        with open(filePath, "r", encoding="utf-8") as f:
            lines = f.readlines()

        with open(filePath, "w", encoding="utf-8") as f_w:
            for line in lines:
                for word in wordlist:
                    if word in line:
                        line = line.replace(word, "*")  # 替换为*
                    f_w.write(line)
        renameAndDelete(filePath, fileName)
    except Exception as e:
        print(e)
        untreatedFile(filePath)


# 处理doc文件
def modifyDoc(filePath, words, fileName):
    wordlist = list(words)
    try:
        doc = Document(filePath)
        for paragraph in doc.paragraphs:
            list_runs = copy.deepcopy(paragraph.runs)
            paragraph.clear()
            for run in list_runs:
                for name in wordlist:
                    if name in run.text:
                        run.text = run.text.replace(name, str("*"))
                # 段落样式的复制
                paragraph.add_run(run.text, run.style)
        doc.save(filePath)
        renameAndDelete(filePath, fileName)
    except Exception as e:
        print(e)
        untreatedFile(filePath)


# 重命名并删除
def renameAndDelete(filePath, fileName):
    path = os.path.split(filePath)[0]
    newFileName = ('.').join(fileName.split('.')[:-1]) + "_待删除"
    suffix = fileName.split('.')[-1]
    os.rename(filePath, path + '\\' + newFileName + suffix)
    deleteFile(path + '\\' + newFileName + suffix)
    print('修改后删除')


# 处理检查到的可疑文件
def dealFile(datas):
    i = 0
    print('总计' + str(len(datas)) + '条可疑记录')
    highDegreeDatas = []
    fileCatagoryList = []
    for data in datas:
        suspectedDegree = data[7]
        filekind = data[13]
        fileCatagoryList.append(filekind)
        if (suspectedDegree == '★ ★ ★ ★ ★ '):  # 遍历json，找出所有★ ★ ★ ★ ★ 以上的文件
            i = i + 1
            highDegreeDatas.append(data)

    fileCatagorySet = set(fileCatagoryList)
    print('所有文件格式：' + str(fileCatagorySet))
    print('★ ★ ★ ★ ★ 可疑文件共计' + str(i) + '条')
    print('开始处理文件')

    for highDegreeData in highDegreeDatas:
        index = highDegreeData[1]
        fileName = highDegreeData[2]
        filePath = highDegreeData[3]
        word = highDegreeData[4]
        content = highDegreeData[5]
        fileSize = highDegreeData[8]
        fileCatagory = highDegreeData[13]

        print('----------------------------------------')
        print('序号：' + str(index))
        print('文件名称：' + str(fileName))
        print('文件路径：' + str(filePath))
        print('关键词：' + str(word))
        print('文本片段：' + str(content))
        print('文件大小：' + str(fileSize) + 'KB')
        print('文件类型：' + str(fileCatagory))
        # 开始处理
        # 可以直接删除的文件有{'JPG 文件', 'WPS PDF 文档','PNG 文件','GIF 文件'}
        allowDeleteFiles = {'JPG 文件', 'WPS PDF 文档', 'PNG 文件', 'GIF 文件'}
        # 不能直接删除的文件有{ '文本文档', 'DOCX 文档', 'XLSX 工作表', 'HTM 文件', 'XLS 工作表', 'PPT 演示文稿', 'HTML 文档', 'PPTX 演示文稿', 'RTF 文件', 'XML 文件', 'DOC 文档'}
        notAllowDeleteFiles = {'文本文档', 'DOCX 文档', 'XLSX 工作表', 'HTM 文件', 'XLS 工作表', 'PPT 演示文稿',
                               'HTML 文档', 'PPTX 演示文稿', 'RTF 文件', 'XML 文件', 'DOC 文档'}
        # 压缩包不处理{'WinRAR ZIP 压缩文件', 'WinRAR 压缩文件管理器'}
        zipFiles = {'WinRAR ZIP 压缩文件', 'WinRAR 压缩文件管理器'}

        if (filePath.count('|') > 0):  # 压缩包不处理
            untreatedFile(filePath)
            continue

        if (not os.path.exists(filePath)):
            print('文件已删除')
            continue

        if (fileCatagory in allowDeleteFiles):
            deleteFile(filePath)
        elif (fileCatagory in notAllowDeleteFiles):
            modifyAndDeleteFile(filePath, word, fileName, fileCatagory)
        else:
            untreatedFile(filePath)


# 解析报告，查询所有记录
with open(htmlPath, 'r', encoding='UTF-8') as f:
    soup = BeautifulSoup(f.read(), 'lxml')
    titles = soup.select('head > script:nth-of-type(2)')
    title = str(titles[0])
    pattern = re.compile("var TableData1={([\s\S]*?)\"};")
    match = re.search(pattern, title, flags=0).group()
    match = match.lstrip('var TableData1=').rstrip(';')
    json = json.loads(match)
    data = json.get('data')
    dealFile(data)  # 开始处理可疑文件
